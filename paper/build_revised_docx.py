from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
RESULTS = ROOT / "experiments" / "results" / "results.json"
OUT = PAPER / "연합학습_그래디언트유출_공격과방어_논문.docx"


FONT = "Arial Unicode MS"
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(85, 85, 85)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def apply_font(run, size: float | None = None, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    apply_font(run, 9, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    r_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attr}"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        apply_font(r1)
        r2 = p.add_run(text[len(bold_prefix):])
        apply_font(r2)
    else:
        r = p.add_run(text)
        apply_font(r)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    apply_font(r, 10.5)


def stat_text(stat: dict, digits: int = 3) -> str:
    return f"{stat['mean']:.{digits}f} ± {stat['std']:.{digits}f}"


def add_results_table(doc: Document, results: dict) -> None:
    attack = results["fedlog_required"]["attack"]
    rows = [
        ("Plain", "plain"),
        ("FedLoG naive", "fedlog_naive"),
        ("FedLoG adaptive", "fedlog_adaptive"),
        ("Noise 0.0001", "noise_0.0001"),
        ("Noise 0.001", "noise_0.001"),
        ("Sparse 90%", "sparse_0.90"),
        ("Sparse 95%", "sparse_0.95"),
    ]
    doc.add_paragraph("표 1. Cora 공식 분할 기반 FedLoG 식 재검증 결과(n=10, 중복 제거)")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["조건", "batch 1 코사인", "batch 1 NMSE", "batch 4 코사인", "batch 4 NMSE"]
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for label, key in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], stat_text(attack["1"][key]["cosine"]))
        set_cell_text(cells[2], stat_text(attack["1"][key]["nmse"], 1))
        set_cell_text(cells[3], stat_text(attack["4"][key]["cosine"]))
        set_cell_text(cells[4], stat_text(attack["4"][key]["nmse"], 1))


def main() -> None:
    results = json.loads(RESULTS.read_text())
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FedLoG 기반 실험 환경에서의 노드 특징 그래디언트 역추론 분석")
    run.bold = True
    apply_font(run, 18, RGBColor(0, 0, 0))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("GNN 연합학습의 제한적 위협 모델과 방어 비교")
    apply_font(r, 11, MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("진로 탐색 프로젝트 · 2026")
    apply_font(r, 10, MUTED)

    doc.add_heading("초록", level=1)
    add_para(
        doc,
        "연합학습은 원본 데이터를 중앙 서버에 보내지 않는다는 점에서 프라이버시 친화적 학습 방식으로 여겨진다. 그러나 클라이언트가 공유한 그래디언트나 모델 업데이트만으로도 원본 입력을 추정할 수 있는 DLG 계열 공격이 알려지며 이러한 가정은 더 이상 충분하지 않다. 본 연구는 FedLoG의 전체 성능 재현이나 정식 프라이버시 보장을 목표로 하지 않고, FedLoG의 2-layer GraphSAGE와 합성 노드/특징 변환 식을 활용한 제한적 실험 환경에서 개별 클라이언트 업데이트가 노드 특징을 얼마나 노출하는지 분석한다. Cora 공식 분할 기반 재검증(n=10)에서는 batch size 1에서 무방어 복원 코사인이 0.754 ± 0.123, FedLoG 특징 변환 포함 조건이 0.773 ± 0.116으로 큰 차이를 보이지 않았다. 반면 90-95% 희소화는 복원 코사인을 0.014-0.028 수준으로 낮췄다.",
    )

    doc.add_heading("1. 서론", level=1)
    add_para(
        doc,
        "DLG 공격은 공유 그래디언트만으로 입력 데이터를 복원할 수 있음을 보였고, 이후 iDLG와 Inverting Gradients는 라벨 추론과 코사인 기반 그래디언트 매칭으로 공격 안정성을 높였다. 그래프 신경망의 연합학습에서는 노드 특징과 그래프 구조, 메시지 패싱이 함께 얽혀 있어 위협 모델을 더 조심스럽게 정의해야 한다.",
    )
    add_bullet(doc, "FedLoG를 프라이버시 방어 모델로 단정하지 않고, 합성 노드와 특징 변환이 유출을 줄이는지 별도 실험으로 검증하였다.")
    add_bullet(doc, "시드 고정이 없는 단일 실행 보고의 위험성을 보이고, Cora 공식 분할 기반 n=10 재검증을 추가하였다.")
    add_bullet(doc, "희소화와 가우시안 노이즈 교란은 기존 방어/압축 기법의 비교 대상으로 적용하되, 정식 DP 보장이나 신규 방어로 과장하지 않았다.")

    doc.add_heading("2. 연구 배경과 해석 범위", level=1)
    add_para(
        doc,
        "전형적인 FedAvg는 클라이언트가 여러 로컬 스텝을 수행한 뒤 모델 파라미터 또는 업데이트량을 서버에 보낸다. 본 연구는 모든 FedAvg 구현이 원시 그래디언트를 전송한다고 가정하지 않는다. 대신 secure aggregation 없이 개별 클라이언트의 단일 로컬 스텝 그래디언트 또는 그와 매우 가까운 업데이트가 관찰되는 제한적 환경을 다룬다.",
    )
    add_para(
        doc,
        "FedLoG는 서브그래프 연합학습에서 local generalization을 개선하기 위한 방법이다. 핵심 목적은 프라이버시 방어가 아니라 로컬 데이터에 부족한 클래스/구조 정보를 합성 노드와 지식 응축으로 보완하는 것이다. 따라서 본 논문은 FedLoG를 프라이버시 방어 모델로 서술하지 않는다.",
    )

    doc.add_heading("3. 실험 방법", level=1)
    add_para(
        doc,
        "공격자는 모델, 손실 함수, 그래프 구조, 타겟 배치의 라벨을 알고 있으며 집계 전 타겟 클라이언트의 업데이트를 관찰한다고 가정한다. 복원 대상은 전체 그래프가 아니라 타겟 노드 또는 작은 배치의 노드 특징이다.",
    )
    add_para(
        doc,
        "공식 FedLoG Cora 분할을 사용했고, 모델은 2-layer GraphSAGE 인코더와 선형 분류기로 구성하였다. 합성 노드는 클래스별 로컬 평균 특징을 class-rate로 가중 집계한 CPU 재현이며, FedLoG 전체 condensation 파이프라인을 완전 재현한 것은 아니다. 공격은 Adam(lr 0.12), 200 iterations, batch size 1과 4에서 수행하였다. 환경은 Python 3.10.20, PyTorch 2.0.1, PyG 2.4.0, CPU 실행이다.",
    )

    doc.add_heading("4. 실험 결과", level=1)
    add_results_table(doc, results)
    doc.add_paragraph()
    doc.add_picture(str(PAPER / "figures" / "fig4_fedlog_required.png"), width=Inches(6.2))
    cap = doc.add_paragraph("그림 1. Cora 공식 분할 기반 FedLoG 식 재검증의 복원 코사인")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "batch size 1에서는 무방어 코사인이 0.754 ± 0.123으로 높았고, FedLoG naive와 adaptive 조건도 각각 0.773 ± 0.116, 0.773 ± 0.112로 거의 같았다. 이는 본 설정에서 합성 특징/특징 변환이 노드 특징 복원을 줄이지 못했음을 의미한다.",
    )
    add_para(
        doc,
        "가우시안 노이즈 0.001은 batch size 4에서 코사인을 0.309 ± 0.087로 낮췄지만, batch size 1에서는 0.734 ± 0.123에 머물렀다. 90%와 95% 희소화는 두 batch size 모두에서 복원 코사인을 거의 0에 가깝게 낮췄다. 다만 이는 기존 압축/방어 기법의 적용 결과이지 새로운 방어 방법의 제안은 아니다.",
    )

    doc.add_heading("5. 논의와 결론", level=1)
    add_para(
        doc,
        "본 실험 범위에서는 FedLoG의 합성 특징과 특징 변환을 포함해도 노드 특징 역추론이 줄어든다고 보기 어려웠다. 이는 FedLoG가 원래 프라이버시 방어 논문이 아니라 local generalization 개선 논문이라는 점과도 일치한다.",
    )
    add_para(
        doc,
        "본 연구는 Cora와 WikiCS 일부 설정, 2-layer GraphSAGE, 작은 batch size, 알려진 구조와 라벨, CPU 재현에 한정된다. 결론은 FedLoG 기반 제한 실험 환경에서 노드 특징 유출 가능성이 확인되었다는 범위로 제한해야 하며, GNN 기반 FL 전체 또는 FedLoG 원 모델 전체의 안전성에 대한 일반 결론으로 확장해서는 안 된다.",
    )

    doc.add_heading("참고문헌", level=1)
    refs = [
        "McMahan et al., Communication-Efficient Learning of Deep Networks from Decentralized Data, AISTATS, 2017.",
        "Zhu, Liu, and Han, Deep Leakage from Gradients, NeurIPS, 2019.",
        "Zhao, Mopuri, and Bilen, iDLG: Improved Deep Leakage from Gradients, arXiv:2001.02610, 2020.",
        "Geiping et al., Inverting Gradients: How Easy Is It to Break Privacy in Federated Learning?, NeurIPS, 2020.",
        "Abadi et al., Deep Learning with Differential Privacy, CCS, 2016.",
        "Kim et al., Subgraph Federated Learning for Local Generalization, ICLR, 2025.",
        "Wei et al., GraphDLG: Exploring Deep Leakage from Gradients in Federated Graph Learning, arXiv:2601.19745, 2026.",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
